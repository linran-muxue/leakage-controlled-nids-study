from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MATERIALS = ROOT / "results_paper_materials_v2"
SOURCE = MATERIALS / "full_paper_body_v3.md"
OUTPUT = ROOT / "论文完整正文_v4_期刊优化稿.docx"


def set_run_font(run, chinese="宋体", western="Times New Roman", size=10.5, bold=None):
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_inline_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, chinese="等线", western="Consolas", size=9.5)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_markdown_body(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_references = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            add_inline_text(p, line[2:])
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            set_run_font(run, chinese="黑体", size=16, bold=True)
            continue
        if line.startswith("## "):
            heading = line[3:]
            in_references = heading == "参考文献"
            p = doc.add_heading(heading, level=1)
            for run in p.runs:
                set_run_font(run, chinese="黑体", size=12, bold=True)
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                set_run_font(run, chinese="黑体", size=11, bold=True)
            continue
        if line.startswith("$$") and line.endswith("$$"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:-2])
            set_run_font(run, chinese="Cambria Math", western="Cambria Math", size=10.5)
            continue
        if re.match(r"^（\d+）", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_text(p, re.sub(r"^（\d+）", "", line).strip())
            continue
        p = doc.add_paragraph()
        if not in_references:
            p.paragraph_format.first_line_indent = Pt(21)
        add_inline_text(p, line)


def add_table(doc, path, title, max_rows=30):
    if not path.exists():
        return
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, chinese="黑体", size=10.5, bold=True)
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(rows[0]):
        table.rows[0].cells[i].text = value
    for row in rows[1:max_rows + 1]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=8)


def add_figure(doc, path, caption):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, size=9.5)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    add_markdown_body(doc)

    doc.add_heading("附录A 主要实验表与图", level=1)
    table_dir = MATERIALS / "tables"
    add_table(doc, table_dir / "table_data_audit.csv", "表A1 CIC-IDS2017数据审计结果", 5)
    add_table(doc, table_dir / "table_split_class_counts.csv", "表A2 训练、验证和测试集类别分布", 20)
    add_table(doc, table_dir / "table_v2_model_comparison.csv", "表A3 统一配置主模型比较", 12)
    add_table(doc, table_dir / "table_v3_individually_tuned_baselines.csv", "表A4 各模型训练集内独立调参结果", 10)
    add_table(doc, table_dir / "table_v2_rf_feature_ablation.csv", "表A5 全特征与χ²特征随机森林消融", 10)
    add_table(doc, table_dir / "table_v2_repeated_split_summary.csv", "表A6 初步三次独立重复划分结果", 10)
    add_table(doc, table_dir / "table_v2_equal_weight_ablation.csv", "表A7 等权与验证集性能加权消融", 10)
    add_table(doc, table_dir / "table_v2_feature_stability.csv", "表A8 χ²特征集合稳定性", 10)
    add_table(doc, table_dir / "table_v3_repeated_10split_summary.csv", "表A9 10次独立重复划分汇总", 10)
    add_table(doc, table_dir / "table_v3_repeated_10split_tests.csv", "表A10 重复划分配对检验", 10)
    add_table(doc, table_dir / "table_v3_weight_mechanism.csv", "表A11 加权机制权重与概率差异分析", 10)
    add_table(doc, table_dir / "table_v2_class_metrics_rf_chi2.csv", "表A12 CIC-IDS2017类别级指标", 15)
    add_table(doc, table_dir / "table_v2_confusion_matrix_rf_chi2.csv", "表A13 CIC-IDS2017混淆矩阵", 10)
    add_table(doc, table_dir / "table_nsl_kdd_results.csv", "表A14 NSL-KDD公平基准", 10)
    add_table(doc, table_dir / "table_nsl_kdd_class_metrics_extra_trees.csv", "表A15 NSL-KDD类别级指标", 15)
    add_table(doc, table_dir / "table_nsl_kdd_confusion_matrix_extra_trees.csv", "表A16 NSL-KDD混淆矩阵", 10)
    add_table(doc, table_dir / "table_v3_imbalanced_class_metrics_rf_chi2.csv", "表A17 自然分布类别级指标", 15)
    add_table(doc, table_dir / "table_v3_imbalanced_confusion_matrix_rf_chi2.csv", "表A18 自然分布混淆矩阵", 10)
    add_table(doc, table_dir / "table_v4_feature_selection_probability_metrics.csv", "表A19 特征选择方法与概率质量比较", 10)
    add_table(doc, table_dir / "table_v4_deployment_latency_percentiles.csv", "表A20 单线程/多线程P50、P95与P99延迟", 25)
    add_table(doc, table_dir / "table_v4_shared_robustness_summary.csv", "表A21 共享扰动鲁棒性比较", 20)
    add_table(doc, table_dir / "table_v4_file_label_coverage.csv", "表A22 原始文件与攻击标签覆盖统计", 12)
    add_table(doc, table_dir / "table_v4_file_label_matrix.csv", "表A23 原始文件—标签数量矩阵", 12)
    add_table(doc, table_dir / "table_v4_nsl_kdd_class_counts.csv", "表A24 NSL-KDD官方划分类别计数", 12)
    add_table(doc, table_dir / "table_v4_nsl_kdd_minority_analysis.csv", "表A25 NSL-KDD R2L/U2R少数类分析", 20)
    add_table(doc, table_dir / "table_v4_feature_quality_summary.csv", "表A26 训练集特征质量审计汇总", 25)
    add_table(doc, table_dir / "table_v4_high_correlation_pairs.csv", "表A27 高相关特征对", 20)
    add_table(doc, table_dir / "table_v4_data_provenance.csv", "表A28 数据来源与版本协议", 10)
    add_table(doc, table_dir / "table_v4_normalized_confusion_matrix_rf_chi2.csv", "表A29 χ²随机森林归一化混淆矩阵", 10)
    add_table(doc, table_dir / "table_v4_statistical_effects_holm.csv", "表A30 配对检验效应量与Holm校正", 15)
    add_table(doc, table_dir / "table_v4_probability_metric_bootstrap_ci.csv", "表A31 概率指标Bootstrap置信区间", 20)
    add_table(doc, table_dir / "table_v4_paired_macro_f1_bootstrap_ci.csv", "表A32 Macro-F1成对Bootstrap差值", 15)
    add_table(doc, table_dir / "table_v4_feature_selection_frequency.csv", "表A33 χ²特征入选频次", 25)
    add_table(doc, table_dir / "table_v4_protocol_sensitivity.csv", "表A34 去重协议敏感性实验", 20)
    add_table(doc, table_dir / "table_v4_imbalanced_predicted_class_counts.csv", "表A35 自然分布测试集类别预测数量", 15)

    figure_dir = MATERIALS / "figures"
    add_figure(doc, figure_dir / "fig_v2_model_performance.png", "图A1 统一配置模型性能比较")
    add_figure(doc, figure_dir / "fig_v2_feature_stability.png", "图A2 χ²特征集合稳定性")
    add_figure(doc, figure_dir / "fig_v2_latency.png", "图A3 离线核心推理延迟")
    add_figure(doc, figure_dir / "fig_v2_robustness.png", "图A4 扰动敏感性")
    add_figure(doc, figure_dir / "fig_v4_calibration_curves.png", "图A5 测试集概率校准曲线")
    add_figure(doc, figure_dir / "fig_v4_latency_p50.png", "图A6 单线程/多线程P50延迟")
    add_figure(doc, figure_dir / "fig_v4_feature_selection_frequency.png", "图A7 χ²特征入选频次")

    doc.save(OUTPUT)
    print(f"COMPLETE_DOCX_WRITTEN={OUTPUT}")


if __name__ == "__main__":
    main()

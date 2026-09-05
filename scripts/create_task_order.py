from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"C:\Users\27677\Documents\ChatGPT\论文\CIC-IDS2017入侵检测论文项目任务单.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color=(255, 255, 255))
        set_cell_shading(hdr.cells[i], '1F4E78')
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_task(doc, title, owner, goal, inputs, actions, outputs, acceptance):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run('负责人：').bold = True
    p.add_run(owner)
    p = doc.add_paragraph()
    p.add_run('任务目标：').bold = True
    p.add_run(goal)
    p = doc.add_paragraph()
    p.add_run('输入：').bold = True
    p.add_run(inputs)
    p = doc.add_paragraph()
    p.add_run('操作步骤：').bold = True
    add_numbered(doc, actions)
    p = doc.add_paragraph()
    p.add_run('预期输出：').bold = True
    add_bullets(doc, outputs)
    p = doc.add_paragraph()
    p.add_run('验收标准：').bold = True
    add_bullets(doc, acceptance)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

styles = doc.styles
styles['Normal'].font.name = 'Microsoft YaHei'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
styles['Normal'].font.size = Pt(10.5)
for name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[name].font.name = 'Microsoft YaHei'
    styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
styles['Heading 1'].font.color.rgb = RGBColor(31, 78, 121)
styles['Heading 2'].font.color.rgb = RGBColor(46, 116, 181)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CIC-IDS2017入侵检测论文项目任务单')
r.bold = True
r.font.size = Pt(20)
r.font.name = 'Microsoft YaHei'
title.paragraph_format.space_after = Pt(6)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('题目：《基于卡方特征优选与加权随机森林的网络入侵检测算法研究》')
rs.font.size = Pt(11)
rs.font.color.rgb = RGBColor(89, 89, 89)
date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date.add_run('版本：实验准备版　|　日期：2026年9月2日').font.size = Pt(9)

doc.add_heading('一、使用说明', level=1)
doc.add_paragraph('本任务单用于指导从数据集准备、算法实现、真实实验到论文定稿的全过程。所有实验数字必须由实际运行产生，不得预先填入或人为修改。你负责电脑上的下载、运行、确认和最终署名；我负责研究设计、代码、数据分析、论文撰写和修改。')
doc.add_paragraph('重要原则：不进行扫网、渗透或攻击；只使用公开数据集；不伪造数据、图表、文献或实验结论；测试集不得参与特征选择、树权重计算或超参数调优。')

doc.add_heading('二、项目目标与研究方案', level=1)
add_table(doc, ['项目项', '确定内容'], [
    ('研究对象', 'CIC-IDS2017机器学习特征数据'),
    ('核心方法', '训练集拟合归一化与卡方特征选择；验证集计算单树表现并进行加权投票'),
    ('主任务', '五分类：Normal、DoS/DDoS、Brute Force、Web Attack、Bot'),
    ('主实验排除标签', 'PortScan、Infiltration、Heartbleed（论文中说明原因）'),
    ('扩展实验', '将上述三个标签归入Other，进行六分类验证（可选但建议完成）'),
    ('主要基线', '决策树、SVM、传统随机森林'),
    ('主要指标', 'Accuracy、Macro-Precision、Macro-Recall、Macro-F1、训练时间、推理延迟'),
    ('重复实验', '至少3个随机种子；建议使用42、2024、3407')
], [4, 12])

doc.add_heading('三、当前已完成状态', level=1)
add_bullets(doc, [
    'Python 3.11.4、pip、Git已确认可用。',
    'E:\\论文项目目录及data、src、results、paper、scripts子目录已创建。',
    '虚拟环境路径为E:\\论文\\.venv\\Scripts\\python.exe。',
    'CIC-IDS2017 MachineLearningCSV已解压到E:\\论文\\data\\raw\\MachineLearningCVE\\。',
    '已确认8个CSV文件、每个79列（78个特征+Label），总样本约283万条。',
    '已完成原始标签统计，数据可用于后续实验。'
])

doc.add_heading('四、文件夹与文件约定', level=1)
doc.add_paragraph('后续所有文件按以下结构保存，不要覆盖raw目录中的原始CSV：')
add_table(doc, ['路径', '用途', '由谁产生'], [
    ('E:\\论文\\data\\raw\\MachineLearningCVE\\', '原始CSV，只读保存', '你已完成'),
    ('E:\\论文\\data\\processed\\', '清洗、映射、划分后的数据', '程序生成'),
    ('E:\\论文\\src\\', '预处理、模型和实验代码', '我编写'),
    ('E:\\论文\\results\\', '指标、混淆矩阵、特征排名、日志', '程序生成'),
    ('E:\\论文\\paper\\', '论文正文、图表、参考文献', '我撰写/整理'),
    ('E:\\论文\\requirements-lock.txt', '固定Python依赖版本', '你运行命令生成')
], [6, 7, 3])

doc.add_heading('五、详细任务单', level=1)

add_task(doc, '任务1：确认研究标签方案', '你确认，我记录', '锁定主实验与扩展实验的标签处理规则，避免根据结果临时改标签。', '已统计的15种原始Label', [
    '确认主实验采用五分类：BENIGN→Normal；DDoS及所有DoS标签→DoS/DDoS；FTP-Patator、SSH-Patator及Web Attack–Brute Force→Brute Force；Web Attack–XSS和Web Attack–Sql Injection→Web Attack；Bot→Bot。',
    '确认主实验排除PortScan、Infiltration、Heartbleed。',
    '确认扩展实验是否将这三个标签统一映射为Other。'
], ['一条明确的确认消息；例如“同意五分类并排除三个标签，扩展实验归入Other”。'], ['映射规则在实验开始前固定；论文方法章节、代码配置和结果表使用同一规则。'])

add_task(doc, '任务2：保存实验环境', '你执行命令，我检查', '固定依赖版本，保证后续实验可以复现。', 'E:\\论文\\.venv', [
    '在PowerShell执行：& "E:\\论文\\.venv\\Scripts\\python.exe" -m pip freeze | Out-File -Encoding utf8 "E:\\论文\\requirements-lock.txt"。',
    '执行：Test-Path "E:\\论文\\.venv\\Scripts\\python.exe"，应返回True。',
    '执行：& "E:\\论文\\.venv\\Scripts\\python.exe" -c "import numpy,pandas,scipy,sklearn,matplotlib,seaborn,joblib,openpyxl; print(\'OK\')"。'
], ['requirements-lock.txt', '终端输出OK'], ['虚拟环境路径存在；所有依赖导入成功；运行脚本时始终使用E:\\论文\\.venv\\Scripts\\python.exe。'])

add_task(doc, '任务3：运行数据集统计与质量检查', '我提供脚本，你运行并回传结果', '确认列名、数据类型、缺失值、无穷值、重复行和标签分布。', '8个原始CSV', [
    '运行我提供的dataset_profile.py，不直接修改原始CSV。',
    '等待程序完成并查看results/dataset_profile.csv、results/label_counts.csv。',
    '把终端输出和两个结果文件发给我，用于核对论文中的数据集描述。'
], ['dataset_profile.csv', 'label_counts.csv', '数据质量日志'], ['所有文件均识别到Label列；异常值统计有记录；标签计数与已知统计基本一致；没有无说明的删除操作。'])

add_task(doc, '任务4：执行数据预处理与标签映射', '我编写程序，你运行', '以分块读取方式完成清洗、标签归并、固定随机种子抽样和分层划分。', '原始CSV、任务1的标签方案', [
    '运行prepare_dataset.py。程序将清理列名、删除NaN和正负无穷样本，并记录删除数量。',
    '程序只在保留标签上进行分层抽样；训练集、验证集、测试集建议按70%/15%/15%划分，具体比例在代码配置中固定。',
    '程序输出映射表、类别计数、划分计数和配置文件。'
], ['data/processed/train.parquet或train.csv', 'validation.parquet或validation.csv', 'test.parquet或test.csv', 'label_mapping.csv', 'dataset_summary.csv', 'preprocess_config.json'], ['三份数据无交集；各类别在划分中有记录；测试集未用于任何拟合步骤；程序可用同一配置重复运行。'])

add_task(doc, '任务5：实现卡方特征优选', '我实现并测试', '在无数据泄漏的前提下筛选20—30个高关联特征，并保存排名。', '训练集、验证集、测试集', [
    '仅使用训练集拟合MinMaxScaler，使特征非负；将同一变换应用于验证集和测试集。',
    '仅在训练集上调用chi2计算特征分数，按预设k值或百分比选择特征。',
    '保存特征名称、chi2分数、p值、排名和最终保留列表。'
], ['results/feature_scores.csv', 'results/selected_features.json'], ['验证集和测试集没有参与特征排名；选中特征数量符合配置；论文中的特征列表与CSV一致。'])

add_task(doc, '任务6：实现基线模型', '我实现，你运行', '建立可比较的决策树、SVM和传统随机森林基线。', '处理后的数据与统一特征矩阵', [
    '使用相同训练/验证/测试划分和相同评价指标。',
    'SVM只在合理规模的训练样本上运行，必要时使用固定分层抽样；记录抽样规则。',
    '保存每个模型的参数、训练时间、预测时间、分类报告和混淆矩阵。'
], ['results/baseline_metrics.csv', '各基线混淆矩阵PNG', 'model_config.json'], ['所有基线使用相同测试集；参数可追溯；结果不使用预设数字。'])

add_task(doc, '任务7：实现加权随机森林', '我实现并测试', '根据验证集上的单树表现计算权重，并进行加权投票。', '随机森林中的决策树、验证集', [
    '训练随机森林后，逐棵树在验证集上计算balanced accuracy或Macro-F1。',
    '将非负树表现归一化：Wi=(score_i+epsilon)/Σ(score_j+epsilon)，记录epsilon。',
    '对测试样本汇总各树对各类别的权重，取得分最高的类别作为最终预测。',
    '同时保留传统多数投票结果，确保能够进行公平对照。'
], ['results/weighted_rf_metrics.csv', 'tree_weights.csv', '加权与普通投票混淆矩阵'], ['树权重只由训练过程中的验证集计算；测试集只在最终评估时使用；权重和约等于1；算法结果可重复。'])

add_task(doc, '任务8：运行完整对比与消融实验', '你运行，我分析', '分别证明特征选择和加权投票的作用。', '任务6、任务7模型代码', [
    '运行三种随机种子，条件允许时运行五种随机种子。',
    '完成四组核心消融：全部特征+普通投票；卡方特征+普通投票；全部特征+加权投票；卡方特征+加权投票。',
    '记录Accuracy、Macro-Precision、Macro-Recall、Macro-F1、训练耗时和推理延迟。',
    '生成均值、标准差、混淆矩阵和特征数量对比。'
], ['results/metrics_summary.csv', 'results/ablation_results.csv', 'results/runtime_summary.csv', 'results/figures/'], ['四组实验均使用相同数据划分；至少3个随机种子；每个结论都有对应数字或图表支撑。'])

add_task(doc, '任务9：实验结果审计', '我完成初审，你完成事实确认', '检查数字一致性、代码与论文一致性和是否存在过度结论。', '全部results文件、运行日志', [
    '我核对CSV中的指标、表格和图的数字是否一致。',
    '我检查摘要、结论中的“显著提升”“速度提升”等表述是否有统计依据。',
    '你确认实验确实在自己的电脑上运行过，数据和结果没有被人为修改。'
], ['results/audit_report.md', '最终可用指标表'], ['论文只写真实结果；任何负结果或提升很小的结果都如实保留并解释。'])

add_task(doc, '任务10：撰写论文正文', '我起草，你核对个人信息和事实', '完成可投稿版本的中文论文、图表和参考文献。', '审计后的实验结果', [
    '先写第三章算法设计和第四章实验分析，再根据实际结果回写摘要、绪论和结论。',
    '补充算法流程图、加权投票伪代码、特征排名图、混淆矩阵和消融表。',
    '统一术语：卡方特征选择、验证集、加权随机森林、Macro-F1等。',
    '参考文献只使用可核验来源，不编造作者、标题、卷期或DOI。'
], ['paper/论文初稿.docx或.md', '图表文件', '参考文献清单'], ['摘要中的数据与第四章一致；每张图表有编号、标题和正文引用；方法描述可由代码复现。'])

add_task(doc, '任务11：格式修改与投稿准备', '我修改，你最终提交', '按学校或目标期刊格式完成定稿和投稿材料。', '论文初稿、学校/期刊模板', [
    '你提供目标期刊名称、格式模板、字数要求、作者和单位信息。',
    '我按模板调整标题、摘要、关键词、章节、图表、参考文献和英文部分。',
    '你使用学校指定查重系统进行检测，并把重复率报告中的问题段落发给我。',
    '我根据查重和模拟审稿意见完成最后修改；你负责登录投稿系统并确认最终提交。'
], ['投稿版Word/PDF', '投稿信（如需要）', '审稿回复模板（如需要）'], ['作者信息和单位真实；引用规范；无虚假实验与文献；符合目标期刊投稿须知。'])

doc.add_heading('六、你需要执行的PowerShell操作清单', level=1)
doc.add_paragraph('以下命令按顺序执行。每一步的输出或报错都应保留，不要只截图最后结果。')
cmds = [
    ('1. 检查虚拟环境', 'Test-Path "E:\\论文\\.venv\\Scripts\\python.exe"'),
    ('2. 保存依赖版本', '& "E:\\论文\\.venv\\Scripts\\python.exe" -m pip freeze | Out-File -Encoding utf8 "E:\\论文\\requirements-lock.txt"'),
    ('3. 检查原始CSV', 'Get-ChildItem -LiteralPath "E:\\论文\\data\\raw\\MachineLearningCVE" -Filter "*.csv" | Select-Object Name,Length'),
    ('4. 运行脚本（脚本生成后）', '& "E:\\论文\\.venv\\Scripts\\python.exe" "E:\\论文\\src\\prepare_dataset.py"'),
    ('5. 查看结果文件', 'Get-ChildItem -LiteralPath "E:\\论文\\results" -Recurse -File | Select-Object FullName,Length'),
]
for label, command in cmds:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(0.5)
    run = code.add_run(command)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(55, 55, 55)

doc.add_heading('七、结果记录规范', level=1)
add_bullets(doc, [
    '每次运行必须记录日期、代码版本、数据文件、随机种子、参数、运行时长和输出目录。',
    '不得只保留最好的一次结果；多随机种子结果应报告均值和标准差。',
    '训练集用于拟合，验证集用于调参和计算树权重，测试集仅用于最终一次或固定次数评估。',
    '如果电脑内存不足，应减少固定抽样规模并记录原因，不得无记录地删除类别或样本。',
    '所有图表的原始数据必须保留为CSV，便于审计和重新绘图。'
])

doc.add_heading('八、最终定稿检查表', level=1)
checks = [
    '□ 标签映射在代码、表格和正文中完全一致。',
    '□ 训练/验证/测试集无交叉，测试集未参与特征选择和权重计算。',
    '□ 所有实验数字均来自真实运行结果。',
    '□ 报告Macro-F1，不只报告Accuracy。',
    '□ 报告类别分布、混淆矩阵和少数类表现。',
    '□ 消融实验能够分别验证卡方筛选和加权投票。',
    '□ 训练时间和推理时间的测量口径已说明。',
    '□ 参考文献真实、可检索、格式统一。',
    '□ 摘要、结果、结论中的数字相互一致。',
    '□ 作者、单位、基金和致谢信息经你本人确认。',
    '□ 已根据目标期刊最新投稿须知完成格式检查。'
]
for c in checks:
    doc.add_paragraph(c)

doc.add_heading('九、当前下一步', level=1)
doc.add_paragraph('你现在只需要回复两项信息：')
add_numbered(doc, [
    '确认是否采用“主实验五分类，排除PortScan、Infiltration、Heartbleed；扩展实验归入Other”的方案。',
    '执行并保存requirements-lock.txt，然后回复“方案已确认，环境版本已保存”。'
])
doc.add_paragraph('收到确认后，我将开始生成dataset_profile.py和prepare_dataset.py，并按本任务单推进。')

doc.core_properties.title = 'CIC-IDS2017入侵检测论文项目任务单'
doc.core_properties.subject = '卡方特征优选与加权随机森林实验项目'
doc.core_properties.author = 'OpenAI协作助手'
doc.save(OUT)
print(OUT)

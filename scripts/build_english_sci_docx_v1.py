from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'results_paper_materials_v3' / 'english_sci_manuscript_v1.md'
OUTPUT = ROOT / 'results_paper_materials_v3' / 'english_sci_manuscript_v1.docx'

def add_text(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p

def add_csv_table(doc, path, title, max_rows=20):
    if not path.exists():
        return
    df = pd.read_csv(path)
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = ["_".join(str(x) for x in c if str(x) != "nan").strip("_") for c in df.columns]
    df = df.head(max_rows)
    doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = f"{value:.4f}" if isinstance(value, float) else str(value)

def add_figure(doc, path, caption):
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(10.5)
    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    for line in lines:
        if not line.strip():
            continue
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('**Keywords:**'):
            p = doc.add_paragraph()
            p.add_run('Keywords: ').bold = True
            p.add_run(line.split('**Keywords:**', 1)[1].strip())
        elif line.startswith('- '):
            add_text(doc, line[2:].strip(), style='List Bullet')
        elif line[:2].isdigit() and line[2:3] == '.':
            add_text(doc, line, style='List Number')
        elif line.startswith('> '):
            add_text(doc, line[2:].strip())
        elif line.startswith('\\[') or line.startswith('\\]') or line.startswith('\\'):
            p = add_text(doc, line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            add_text(doc, line)
    # Insert reproducible evidence objects after the narrative. The Markdown
    # remains the source of truth; this DOCX is a submission-oriented rendering.
    doc.add_page_break()
    doc.add_heading('Evidence tables and figures', level=1)
    add_csv_table(doc, ROOT / 'results_publication_final' / 'metrics_aggregate.csv', 'Table 1. Locked CIC-IDS2017 model comparison')
    add_csv_table(doc, ROOT / 'results_unsw_nb15_independent_v4' / 'metrics_aggregate.csv', 'Table 2. UNSW-NB15 independent benchmark')
    add_csv_table(doc, ROOT / 'results_paper_materials_v3' / 'tables' / 'table_unsw_minority_analysis.csv', 'Table 3. UNSW-NB15 minority-class analysis', max_rows=12)
    figures = [
        ('fig_v2_model_performance.png', 'Figure 1. CIC-IDS2017 model performance under the locked protocol.'),
        ('fig_v4_calibration_curves.png', 'Figure 2. Probability calibration curves for the CIC-IDS2017 models.'),
        ('fig_v4_feature_selection_frequency.png', 'Figure 3. Feature-selection frequency across repeated splits.'),
        ('fig_v4_latency_p50.png', 'Figure 4. P50 deployment latency under the measurement protocol.'),
    ]
    for name, caption in figures:
        add_figure(doc, ROOT / 'results_paper_materials_v2' / 'figures' / name, caption)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f'ENGLISH_DOCX_WRITTEN={OUTPUT}')

if __name__ == '__main__':
    main()

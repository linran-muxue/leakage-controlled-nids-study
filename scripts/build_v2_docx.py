from pathlib import Path
import csv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

ROOT=Path(__file__).resolve().parents[1]
MAT=ROOT/'results_paper_materials_v2'
OUT=ROOT/'论文投稿草稿_v2_统一无泄漏实验结果.docx'

def add_table(doc,path,max_rows=18):
    rows=list(csv.reader(path.open(encoding='utf-8-sig')))
    if not rows:return
    t=doc.add_table(rows=1,cols=len(rows[0])); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,v in enumerate(rows[0]):t.rows[0].cells[i].text=v
    for row in rows[1:max_rows+1]:
        cells=t.add_row().cells
        for i,v in enumerate(row):cells[i].text=v

def add_img(doc,path,caption):
    if path.exists():
        doc.add_picture(str(path),width=Inches(6.1)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        p=doc.add_paragraph(caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER

def main():
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.8); sec.left_margin=Inches(.9); sec.right_margin=Inches(.9)
    doc.styles['Normal'].font.name='宋体'; doc.styles['Normal'].font.size=Pt(10.5)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=title.add_run('基于训练集交叉验证与卡方特征选择的随机森林网络入侵检测研究'); r.bold=True; r.font.size=Pt(16)
    p=doc.add_paragraph('[作者姓名]  [学校/学院/专业]'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph('通信作者：[姓名]；E-mail：[待填写]'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('摘要',1)
    doc.add_paragraph('针对网络入侵检测中类别不均衡、特征冗余和实验流程易发生数据泄漏等问题，本文构建了一套训练集内预处理、交叉验证选参和独立测试评估的可复现实验流程。以CIC-IDS2017为主数据集，经过标签映射、无效值清理、全局去重、跨标签冲突剔除和分层划分后，比较决策树、支持向量机、ExtraTrees、χ²特征随机森林及验证集加权随机森林。统一配置（χ²保留60个特征、100棵树、min_samples_leaf=2）下，χ²随机森林在平衡五分类子集测试集上取得95.97%的Accuracy和95.98%的Macro-F1。重复分层划分显示该结果具有较好的稳定性，但χ²筛选增益有限；等权投票与验证集加权投票逐样本结果一致，差异未达到统计显著水平。进一步在NSL-KDD上进行独立基准测试，并评估离线推理性能和扰动敏感性。结果显示，模型对随机特征屏蔽相对稳定，但对连续值高斯噪声较敏感。')
    doc.add_paragraph('关键词：网络入侵检测；CIC-IDS2017；随机森林；卡方特征选择；交叉验证；数据泄漏审计')
    doc.add_heading('Title, Abstract and Keywords',1)
    doc.add_paragraph('Leakage-Safe Cross-Validation and Chi-Square Feature Selection for Random-Forest Network Intrusion Detection')
    doc.add_paragraph('Abstract: This paper develops a reproducible network intrusion detection workflow with fold-local preprocessing, chi-square feature selection, cross-validated hyperparameter tuning, and an untouched test set. Experiments on CIC-IDS2017 compare decision trees, SVM, ExtraTrees, chi-square random forests, and validation-score-weighted random forests. Under the locked protocol (60 chi-square-selected features, 100 trees, and min_samples_leaf=2), the chi-square random forest achieves 95.97% accuracy and 95.98% macro-F1 on the balanced five-class subset. The weighted voting variant does not yield a statistically significant improvement. An independent NSL-KDD benchmark, offline inference measurements, and perturbation tests further show that random feature masking has limited impact, whereas continuous Gaussian noise causes a substantial performance decrease.')
    doc.add_paragraph('Keywords: network intrusion detection; CIC-IDS2017; random forest; chi-square feature selection; cross-validation; leakage audit')
    body=(MAT/'chapter4_results_draft.md').read_text(encoding='utf-8')
    for line in body.splitlines():
        if not line.strip(): continue
        if line.startswith('# '): continue
        if line.startswith('## '): doc.add_heading(line[3:],1)
        elif line.startswith('### '): doc.add_heading(line[4:],2)
        else: doc.add_paragraph(line)
    doc.add_heading('主要结果表',1)
    add_table(doc,MAT/'tables/table_v2_model_comparison.csv',max_rows=10)
    add_table(doc,MAT/'tables/table_v2_bootstrap_ci.csv',max_rows=20)
    add_table(doc,MAT/'tables/table_v2_paired_tests.csv',max_rows=20)
    add_table(doc,MAT/'tables/table_nsl_kdd_results.csv',max_rows=10)
    add_table(doc,MAT/'tables/table_nsl_kdd_class_metrics_extra_trees.csv',max_rows=12)
    add_table(doc,MAT/'tables/table_nsl_kdd_confusion_matrix_extra_trees.csv',max_rows=8)
    add_table(doc,MAT/'tables/table_unified_imbalanced_sensitivity.csv',max_rows=20)
    add_table(doc,MAT/'tables/table_v2_repeated_split_summary.csv',max_rows=12)
    add_table(doc,MAT/'tables/table_v2_equal_weight_ablation.csv',max_rows=8)
    add_table(doc,MAT/'tables/table_v2_feature_stability.csv',max_rows=8)
    add_table(doc,MAT/'tables/table_v2_class_metrics_rf_chi2.csv',max_rows=12)
    add_table(doc,MAT/'tables/table_v2_confusion_matrix_rf_chi2.csv',max_rows=8)
    doc.add_heading('图表',1)
    add_img(doc,MAT/'figures/fig_v2_model_performance.png','图1 v2统一调参模型性能对比')
    add_img(doc,MAT/'figures/fig_v2_latency.png','图2 v2离线推理延迟')
    add_img(doc,MAT/'figures/fig_v2_robustness.png','图3 v2扰动鲁棒性结果')
    add_img(doc,MAT/'figures/fig_v2_feature_stability.png','图4 训练集内χ²特征集合稳定性')
    doc.add_heading('结论',1)
    doc.add_paragraph('本文形成了训练集内缩放、折内χ²评分、交叉验证选参、独立测试及数据审计相结合的网络入侵检测实验流程。CIC-IDS2017平衡五分类研究子集上，χ²60随机森林获得95.97%的平均Accuracy和95.98%的平均Macro-F1；重复分层划分实验显示模型平均Macro-F1为95.44%±1.25%，但相对全特征RF的改进幅度较小。等权投票与验证集性能加权投票在当前实验中给出相同预测，统计检验亦未支持加权机制具有显著优势。NSL-KDD官方划分上的独立基准进一步揭示R2L、U2R等少数类仍是主要困难。后续工作应优先开展覆盖完整类别的时间/文件外验证，研究面向极端不平衡类别的代价敏感学习，并在真实流量特征提取链路上评估端到端时延。')
    doc.add_heading('参考文献',1)
    references = [
        '[1] SHARAFALDIN I, LASHKARI A H, GHORBANI A A. Toward generating a new intrusion detection dataset and intrusion traffic characterization[C]//Proceedings of the 4th International Conference on Information Systems Security and Privacy. 2018: 108-116. DOI:10.5220/0006639801080116.',
        '[2] TAVALLAEE M, BAGHERI E, LU W, et al. A detailed analysis of the KDD CUP 99 data set[C]//2009 IEEE Symposium on Computational Intelligence for Security and Defense Applications. 2009: 1-6. DOI:10.1109/CISDA.2009.5356528.',
        '[3] BREIMAN L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32. DOI:10.1023/A:1010933404324.',
        '[4] LIU H, SETIONO R. Chi2: Feature selection and discretization of numeric attributes[C]//Proceedings of the 7th International Conference on Tools with Artificial Intelligence. 1995: 388-391. DOI:10.1109/TAI.1995.479783.',
        '[5] BUCZAK A L, GUVEN E. A survey of data mining and machine learning methods for cyber security intrusion detection[J]. IEEE Communications Surveys & Tutorials, 2016, 18(2): 1153-1176. DOI:10.1109/COMST.2015.2494502.',
        '[6] RING M, WUNDERLICH S, GRUDL D, et al. A survey of network-based intrusion detection data sets[J]. Computers & Security, 2019, 86: 147-167. DOI:10.1016/j.cose.2019.06.005.',
        '[7] KHRAISAT A, GONDAL I, VAMPLEW P, et al. Survey of intrusion detection systems: techniques, datasets and challenges[J]. Cybersecurity, 2019, 2: 20. DOI:10.1186/s42400-019-0038-7.',
        '[8] MCNEMAR Q. Note on the sampling error of the difference between correlated proportions or percentages[J]. Psychometrika, 1947, 12: 153-157. DOI:10.1007/BF02295996.',
        '[9] EFRON B, TIBSHIRANI R J. Improvements on cross-validation: the .632+ bootstrap method[J]. Journal of the American Statistical Association, 1997, 92(438): 548-560. DOI:10.2307/2965703.',
        '[10] 郭健忠, 王灿, 谢斌, 等. 面向车联网DoS攻击的混合入侵检测系统[J]. 计算机系统应用, 2025, 34(3): 85-93. DOI:10.15888/j.cnki.csa.009821.',
        '[11] 刘小龙, 邢东辉, 蔡婷婷, 等. 基于改进差分进化算法的堆叠自编码器网络入侵检测[J]. 计算机系统应用, 2025, 34(11): 95-106. DOI:10.15888/j.cnki.csa.009976.',
        '[12] 谭书志, 周传, 黄世元. 基于深度学习的入侵检测方法[J]. 计算机系统应用, 2025, 34(9): 170-179. DOI:10.15888/j.cnki.csa.009953.',
        '[13] 李贵玲, 龚琼, 邓华山, 等. 融合GRU与CNN的网络入侵检测模型[J]. 计算机系统应用, 2023, 32(8): 162-170. DOI:10.15888/j.cnki.csa.009194.'
    ]
    for ref in references: doc.add_paragraph(ref)
    doc.add_paragraph('注：上述条目已由论文/期刊页面和DOI候选记录核验；正式投稿前仍须按《计算机系统应用》最新模板复核作者截断规则、英文大小写及标点，并在完整第一至三章中建立逐条正文标引。')
    doc.add_heading('作者与投稿信息（待填写）',1)
    doc.add_paragraph('作者、单位、通信地址、邮编、联系电话、E-mail、基金项目和作者简介须由作者提供真实信息。本文档尚未套用《计算机系统应用》官方Word模板。')
    doc.save(OUT); print(f'DOCX_WRITTEN={OUT}')

if __name__=='__main__': main()

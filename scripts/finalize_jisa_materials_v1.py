from pathlib import Path
from docx import Document
from docx.shared import Pt

root = Path(__file__).resolve().parents[1]
md = root / 'results_paper_materials_v3' / 'english_sci_manuscript_v1.md'
text = md.read_text(encoding='utf-8')

# JISA allows 1--7 keywords.
text = text.replace(
    '**Keywords:** network intrusion detection; CIC-IDS2017; NSL-KDD; UNSW-NB15; leakage control; feature selection; random forest; reproducibility; open-set recognition; dataset audit',
    '**Keywords:** network intrusion detection; CIC-IDS2017; feature selection; random forest; leakage control; reproducibility; open-set recognition'
)
text = text.replace(
    'Following the closed-set/open-set distinction in the open-set recognition literature [12],',
    'Following the closed-set/open-set distinction in the open-set recognition literature (Geng et al., 2021),'
)

# Replace numbered reference block with an alphabetized author--year list.
ref_start = text.index('## References\n')
prefix = text[:ref_start]
refs = '''## References

B reiman L. Random forests. Machine Learning, 2001, 45(1):5–32. DOI:10.1023/A:1010933404324.
Buczak AL, Guven E. A survey of data mining and machine learning methods for cyber security intrusion detection. IEEE Communications Surveys & Tutorials, 2016, 18(2):1153–1176. DOI:10.1109/COMST.2015.2494502.
Chen T, Guestrin C. XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016:785–794. DOI:10.1145/2939672.2939785.
Demšar J. Statistical comparisons of classifiers over multiple data sets. Journal of Machine Learning Research, 2006, 7:1–30.
Efron B, Tibshirani RJ. Improvements on cross-validation: The .632+ bootstrap method. Journal of the American Statistical Association, 1997, 92(438):548–560. DOI:10.2307/2965703.
Geng C, Huang SH, Chen S. Recent advances in open set recognition: A survey. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2021, 43(10):3614–3631. DOI:10.1109/TPAMI.2020.2981604.
Geurts P, Ernst D, Wehenkel L. Extremely randomized trees. Machine Learning, 2006, 63:3–42. DOI:10.1007/s10994-006-6226-1.
Khraisat A, Gondal I, Vamplew P, Kamruzzaman J. Survey of intrusion detection systems: Techniques, datasets and challenges. Cybersecurity, 2019, 2:20. DOI:10.1186/s42400-019-0038-7.
Liu H, Setiono R. Chi2: Feature selection and discretization of numeric attributes. ICTAI, 1995, 388–391. DOI:10.1109/TAI.1995.479783.
Lu J, Liu A, Dong F, Gu F, Gama J, Zhang G. Learning under concept drift: A review. IEEE Transactions on Knowledge and Data Engineering, 2019, 31(12):2346–2363. DOI:10.1109/TKDE.2018.2876857.
McNemar Q. Note on the sampling error of the difference between correlated proportions or percentages. Psychometrika, 1947, 12:153–157. DOI:10.1007/BF02295996.
Moustafa N, Slay J. UNSW-NB15: A comprehensive data set for network intrusion detection systems. MilCIS, 2015:1–6. DOI:10.1109/MilCIS.2015.7348942.
Ring M, Wunderlich S, Grüdl D, Landes D, Hotho A. A survey of network-based intrusion detection data sets. Computers & Security, 2019, 86:147–167. DOI:10.1016/j.cose.2019.06.005.
Sharafaldin I, Lashkari AH, Ghorbani AA. Toward generating a new intrusion detection dataset and intrusion traffic characterization. ICISSP, 2018:108–116. DOI:10.5220/0006639801080116.
Tavallaee M, Bagheri E, Lu W, Ghorbani AA. A detailed analysis of the KDD CUP 99 data set. CISDA, 2009:1–6. DOI:10.1109/CISDA.2009.5356528.
'''.replace('B reiman', 'Breiman')
md.write_text(prefix + refs, encoding='utf-8')

# JISA Highlights: each bullet is well below 85 characters.
highlights = [
    'Auditable preprocessing separates curation choices from fitted transforms.',
    'Chi-square filtering reduces features with nearly unchanged macro-F1.',
    'Tree weighting matches equal voting under the locked CIC protocol.',
    'Independent datasets reveal strong class-prior and split effects.',
]
hdir = root / 'results_paper_materials_v3'
(hdir / 'Highlights_JISA.txt').write_text('\n'.join('- ' + x for x in highlights) + '\n', encoding='utf-8')
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(11)
doc.add_heading('Highlights', level=1)
for x in highlights:
    doc.add_paragraph(x, style='List Bullet')
doc.save(hdir / 'Highlights_JISA.docx')
print(f'UPDATED_MANUSCRIPT={md}')
print(f'HIGHLIGHTS={hdir / "Highlights_JISA.docx"}')

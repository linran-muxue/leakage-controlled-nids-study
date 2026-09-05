from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MATERIALS = ROOT / "results_paper_materials"
OUT = ROOT / "论文草稿_基于卡方特征选择与随机森林的CIC-IDS2017入侵检测研究.docx"


def add_csv_table(doc, path, max_rows=25):
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        table.rows[0].cells[i].text = value
    for row in rows[1:max_rows + 1]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph(f"数据来源：{path.name}").style = "Caption"


def add_image(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.2))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 2"].font.name = "黑体"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于卡方特征选择与随机森林的CIC-IDS2017入侵检测研究")
    run.bold = True
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph("论文草稿（基于真实实验结果生成）")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    body = (MATERIALS / "paper_body_draft.md").read_text(encoding="utf-8")
    # Skip title and reference note is included as a final section.
    for line in body.splitlines()[2:]:
        if not line.strip():
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("**关键词：**"):
            doc.add_paragraph(line.replace("**", ""))
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*"))
        elif line.startswith("$$") or line.startswith("其中ε"):
            doc.add_paragraph(line.replace("$$", ""))
        else:
            doc.add_paragraph(re.sub(r"`([^`]*)`", r"\1", line))

    doc.add_page_break()
    doc.add_heading("附录A 实验结果表", level=1)
    add_csv_table(doc, MATERIALS / "tables" / "table_model_comparison.csv", max_rows=10)
    doc.add_heading("附录B 数据审计与划分统计", level=1)
    add_csv_table(doc, MATERIALS / "tables" / "table_data_audit.csv", max_rows=5)
    add_csv_table(doc, MATERIALS / "tables" / "table_split_class_counts.csv", max_rows=20)
    doc.add_heading("附录C 图表", level=1)
    figs = [
        ("fig_chi2_top20.png", "图C-1 χ²值排名前20的特征"),
        ("fig_k_selection.png", "图C-2 特征数量选择曲线"),
        ("fig_model_performance.png", "图C-3 模型性能对比"),
        ("fig_time_comparison.png", "图C-4 训练与推理时间对比"),
        ("fig_confusion_matrix_rf_chi2_seed42.png", "图C-5 随机森林+χ²60混淆矩阵（seed=42）"),
    ]
    for name, caption in figs:
        add_image(doc, MATERIALS / "figures" / name, caption)

    doc.add_heading("附录D 使用说明", level=1)
    doc.add_paragraph("本草稿中的结果来自CIC-IDS2017公开数据集和项目目录中的可复现实验脚本。正式提交前，需要根据学校模板调整封面、目录、图表编号、公式格式和参考文献格式，并由作者核对实验环境、作者信息及引用文献。")
    doc.save(OUT)
    print(f"DOCX_WRITTEN={OUT}")


if __name__ == "__main__":
    main()

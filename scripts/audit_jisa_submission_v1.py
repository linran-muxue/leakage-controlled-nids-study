from pathlib import Path
import re
from docx import Document

root = Path(__file__).resolve().parents[1]
md = root / 'results_paper_materials_v3' / 'english_sci_manuscript_v1.md'
text = md.read_text(encoding='utf-8')
errors = []

abstract = text.split('## 1. Introduction', 1)[0]
abstract_words = len(re.findall(r"\b[\w'-]+\b", abstract.split('## Abstract', 1)[-1]))
if abstract_words > 250:
    errors.append(f'abstract exceeds 250 words: {abstract_words}')
if '[1]' in text or '[2]' in text or '[3]' in text or '[4]' in text or '[5]' in text or '[6]' in text or '[7]' in text or '[8]' in text or '[9]' in text or '[10]' in text or '[11]' in text or '[12]' in text or '[13]' in text or '[14]' in text or '[15]' in text:
    errors.append('numeric citation remains')
kw_line = next((x for x in text.splitlines() if x.startswith('**Keywords:**')), '')
if len([x.strip() for x in kw_line.split('**Keywords:**',1)[-1].split(';') if x.strip()]) > 7:
    errors.append('more than 7 keywords')
for required in ['## Data and code availability', '## Declaration of Generative AI', '### CRediT authorship contribution statement', '### Funding', '### Declaration of competing interest']:
    if required not in text:
        errors.append(f'missing section: {required}')
for f in [root/'results_paper_materials_v3/english_sci_manuscript_v1.docx', root/'results_paper_materials_v3/Highlights_JISA.docx']:
    if not f.exists():
        errors.append(f'missing file: {f.name}')
if errors:
    raise SystemExit('JISA_AUDIT_FAIL: ' + '; '.join(errors))
doc = Document(root/'results_paper_materials_v3/english_sci_manuscript_v1.docx')
print(f'JISA_AUDIT_OK abstract_words={abstract_words} keywords=7 paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}')

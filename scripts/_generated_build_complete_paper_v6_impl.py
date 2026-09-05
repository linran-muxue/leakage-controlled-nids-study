"""Generate the complete manuscript DOCX with the revised UNSW-NB15 section."""
from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "results_paper_materials_v3" / "full_paper_body_v6_data_processing.md"
MATERIALS = ROOT / "results_paper_materials_v2"
UNSW_TABLES = ROOT / "results_paper_materials_v3" / "tables"
OUTPUT = ROOT / "论文完整正文_v6_数据处理完善稿.docx"


def set_run_font(run, chinese="宋体", western="Times New Roman", size=10.5, bold=None):
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_inline_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); set_run_font(run, chinese="等线", western="Consolas", size=9.5)
        else:
            run = paragraph.add_run(part); set_run_font(run)


def add_markdown_body(doc):
    in_references = False
    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(); add_inline_text(p, line[2:]); continue
        if line.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:]); set_run_font(run, chinese="黑体", size=16, bold=True); continue
        if line.startswith("## "):
            heading = line[3:]; in_references = heading == "参考文献"
            p = doc.add_heading(heading, level=1)
            for run in p.runs: set_run_font(run, chinese="黑体", size=12, bold=True)
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs: set_run_font(run, chinese="黑体", size=11, bold=True)
            continue
        if line.startswith("$$") and line.endswith("$$"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:-2]); set_run_font(run, chinese="Cambria Math", western="Cambria Math"); continue
        if re.match(r"^（\d+）", line):
            p = doc.add_paragraph(style="List Number"); add_inline_text(p, re.sub(r"^（\d+）", "", line).strip()); continue
        p = doc.add_paragraph()
        if not in_references: p.paragraph_format.first_line_indent = Pt(21)
        add_inline_text(p, line)


def add_table(doc, path, title, max_rows=30):
    if not path.exists(): return
    p = doc.add_paragraph(title); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs: set_run_font(run, chinese="黑体", size=10.5, bold=True)
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    if not rows: return
    table = doc.add_table(rows=1, cols=len(rows[0])); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(rows[0]): table.rows[0].cells[i].text = value
    for row in rows[1:max_rows + 1]:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs: set_run_font(run, size=8)


def main():
    doc = Document(); section = doc.sections[0]
    section.top_margin = Inches(0.75); section.bottom_margin = Inches(0.75); section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]; normal.font.name = "Times New Roman"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); normal.font.size = Pt(10.5)
    add_markdown_body(doc)
    doc.add_heading("附录A 主要实验表与图", level=1)
    table_dir = MATERIALS / "tables"
    legacy = [
        ("table_data_audit.csv", "表A1 CIC-IDS2017数据审计结果", 5),
        ("table_split_class_counts.csv", "表A2 训练、验证和测试集类别分布", 20),
        ("table_v2_model_comparison.csv", "表A3 统一配置主模型比较", 12),
        ("table_v3_individually_tuned_baselines.csv", "表A4 各模型训练集内独立调参结果", 10),
        ("table_v2_rf_feature_ablation.csv", "表A5 全特征与χ²特征随机森林消融", 10),
        ("table_v3_repeated_10split_summary.csv", "表A6 10次独立重复划分汇总", 10),
        ("table_v3_repeated_10split_tests.csv", "表A7 重复划分配对检验", 10),
        ("table_v3_weight_mechanism.csv", "表A8 加权机制权重与概率差异分析", 10),
        ("table_v2_class_metrics_rf_chi2.csv", "表A9 CIC-IDS2017类别级指标", 15),
        ("table_v2_confusion_matrix_rf_chi2.csv", "表A10 CIC-IDS2017混淆矩阵", 10),
        ("table_v4_feature_selection_probability_metrics.csv", "表A11 特征选择方法与概率质量比较", 10),
        ("table_v4_probability_metric_bootstrap_ci.csv", "表A12 CIC-IDS2017概率指标Bootstrap区间", 20),
        ("table_v4_deployment_latency_percentiles.csv", "表A13 单线程/多线程延迟", 25),
        ("table_v4_shared_robustness_summary.csv", "表A14 共享扰动鲁棒性比较", 20),
        ("table_v4_file_label_matrix.csv", "表A15 原始文件—标签数量矩阵", 12),
        ("table_v4_nsl_kdd_class_counts.csv", "表A16 NSL-KDD官方划分类别计数", 12),
        ("table_v4_nsl_kdd_minority_analysis.csv", "表A17 NSL-KDD R2L/U2R少数类分析", 20),
    ]
    for name, title, rows in legacy: add_table(doc, table_dir / name, title, rows)
    add_table(doc, UNSW_TABLES / "table_v6_processing_stage_counts.csv", "表A29 数据处理阶段计数", 12)
    add_table(doc, UNSW_TABLES / "table_v6_raw_file_processing_audit.csv", "表A30 原始文件逐文件处理审计", 12)
    add_table(doc, UNSW_TABLES / "table_v6_feature_quality_and_leakage_checks.csv", "表A31 特征质量与泄漏检查", 12)
    add_table(doc, UNSW_TABLES / "table_v6_file_external_generalization.csv", "表A32 文件外泛化与类别覆盖审计", 20)
    add_table(doc, UNSW_TABLES / "table_v6_unsw_cross_split_sensitivity.csv", "表A33 UNSW-NB15跨split重复敏感性", 10)
    add_table(doc, UNSW_TABLES / "table_v6_open_set_matrix.csv", "表A34 开放集未知攻击拒识组合", 10)
    unsw = [
        ("table_unsw_provenance.csv", "表A18 UNSW-NB15数据来源与哈希", 5),
        ("table_unsw_metrics_aggregate.csv", "表A19 UNSW-NB15三种子汇总指标", 10),
        ("table_unsw_class_metrics_rf_all_3seeds.csv", "表A20 UNSW-NB15 RF类别级指标", 15),
        ("table_unsw_class_metrics_xgboost_chi2_3seeds.csv", "表A21 UNSW-NB15 XGBoost类别级指标", 15),
        ("table_unsw_confusion_matrix_rf_all_seed2024.csv", "表A22 UNSW-NB15 RF混淆矩阵", 12),
        ("table_unsw_normalized_confusion_matrix_rf_all_seed2024.csv", "表A23 UNSW-NB15 RF归一化混淆矩阵", 12),
        ("table_unsw_confusion_matrix_xgboost_chi2_seed2024.csv", "表A24 UNSW-NB15 XGBoost混淆矩阵", 12),
        ("table_unsw_normalized_confusion_matrix_xgboost_chi2_seed2024.csv", "表A25 UNSW-NB15 XGBoost归一化混淆矩阵", 12),
        ("table_unsw_prediction_counts_rf_all_seed2024.csv", "表A26 UNSW-NB15 RF真实/预测数量", 15),
        ("table_unsw_prediction_counts_xgboost_chi2_seed2024.csv", "表A27 UNSW-NB15 XGBoost真实/预测数量", 15),
        ("table_unsw_minority_analysis.csv", "表A28 UNSW-NB15少数类分析", 15),
    ]
    for name, title, rows in unsw: add_table(doc, UNSW_TABLES / name, title, rows)
    figures = [("fig_v2_model_performance.png", "图A1 统一配置模型性能比较"), ("fig_v4_calibration_curves.png", "图A2 测试集概率校准曲线"), ("fig_v4_feature_selection_frequency.png", "图A3 χ²特征入选频次")]
    for name, caption in figures:
        path = MATERIALS / "figures" / name
        if path.exists():
            doc.add_picture(str(path), width=Inches(5.9)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph(caption); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: set_run_font(run, size=9.5)
    doc.save(OUTPUT); print(f"COMPLETE_DOCX_WRITTEN={OUTPUT}")


if __name__ == "__main__": main()
